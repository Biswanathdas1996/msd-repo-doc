import io
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from fpdf import FPDF


BRAND_NAVY = RGBColor(0x1B, 0x2A, 0x4A)
BRAND_BLUE = RGBColor(0x2B, 0x57, 0x9A)
BRAND_LIGHT = RGBColor(0xE8, 0xEE, 0xF6)
BRAND_ACCENT = RGBColor(0xD4, 0x7B, 0x2E)
COLOR_DARK = RGBColor(0x2D, 0x2D, 0x2D)
COLOR_BODY = RGBColor(0x3A, 0x3A, 0x3A)
COLOR_MUTED = RGBColor(0x6B, 0x6B, 0x6B)
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def _strip_markdown(text: str) -> list[dict]:
    lines = text.split("\n")
    blocks = []
    in_code_block = False
    code_lines = []

    for line in lines:
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code_block:
                in_code_block = False
                blocks.append({"type": "code_block", "text": "\n".join(code_lines)})
                code_lines = []
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        if not stripped:
            blocks.append({"type": "blank", "text": ""})
            continue

        if stripped.startswith("######"):
            blocks.append({"type": "h6", "text": stripped[6:].strip()})
        elif stripped.startswith("#####"):
            blocks.append({"type": "h5", "text": stripped[5:].strip()})
        elif stripped.startswith("####"):
            blocks.append({"type": "h4", "text": stripped[4:].strip()})
        elif stripped.startswith("###"):
            blocks.append({"type": "h3", "text": stripped[3:].strip()})
        elif stripped.startswith("##"):
            blocks.append({"type": "h2", "text": stripped[2:].strip()})
        elif stripped.startswith("#"):
            blocks.append({"type": "h1", "text": stripped[1:].strip()})
        elif stripped.startswith("- ") or stripped.startswith("* "):
            blocks.append({"type": "bullet", "text": stripped[2:].strip()})
        elif re.match(r"^\d+\.\s", stripped):
            blocks.append({"type": "numbered", "text": re.sub(r"^\d+\.\s", "", stripped)})
        elif stripped.startswith("|") and stripped.endswith("|"):
            blocks.append({"type": "table_row", "text": stripped})
        elif stripped.startswith("---") or stripped.startswith("***"):
            blocks.append({"type": "hr", "text": ""})
        else:
            blocks.append({"type": "paragraph", "text": stripped})

    if in_code_block and code_lines:
        blocks.append({"type": "code_block", "text": "\n".join(code_lines)})

    return blocks


def _clean_inline_md(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
    return text


def _set_cell_shading(cell, color_hex: str):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_border(cell, border_color="CCCCCC"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)


def _set_cell_margins(cell, top=40, bottom=40, left=80, right=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def _format_cell_text(cell, text, font_name="Calibri", font_size=9, bold=False, color=None):
    for p in cell.paragraphs:
        p.clear()
    p = cell.paragraphs[0]
    p.space_before = Pt(1)
    p.space_after = Pt(1)
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def _setup_styles(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    font.color.rgb = COLOR_BODY

    pf = style.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(4)
    pf.line_spacing = 1.15

    for level, (size, color) in enumerate([
        (22, BRAND_NAVY), (16, BRAND_NAVY), (13, BRAND_BLUE),
        (11, BRAND_BLUE), (10, COLOR_DARK), (10, COLOR_DARK),
    ], start=1):
        h_style = doc.styles[f"Heading {level}"]
        h_style.font.name = "Calibri"
        h_style.font.size = Pt(size)
        h_style.font.color.rgb = color
        h_style.font.bold = True
        h_style.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
        h_style.paragraph_format.space_after = Pt(4)
        h_style.paragraph_format.keep_with_next = True


def _add_cover_page(doc, solution_name: str):
    for _ in range(6):
        doc.add_paragraph("")

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(solution_name)
    run.font.name = "Calibri"
    run.font.size = Pt(32)
    run.font.color.rgb = BRAND_NAVY
    run.bold = True

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_p.add_run("Technical Design Document")
    run.font.name = "Calibri"
    run.font.size = Pt(18)
    run.font.color.rgb = BRAND_BLUE

    doc.add_paragraph("")

    line_p = doc.add_paragraph()
    line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line_p.add_run("━" * 40)
    run.font.color.rgb = BRAND_ACCENT
    run.font.size = Pt(12)

    doc.add_paragraph("")

    info_items = [
        ("Document Type", "Technical Design Document"),
        ("Solution", solution_name),
        ("Generated", datetime.now().strftime("%B %d, %Y")),
        ("Status", "Draft"),
    ]

    info_table = doc.add_table(rows=len(info_items), cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(info_items):
        row = info_table.rows[i]
        _format_cell_text(row.cells[0], label, font_size=10, bold=True, color=COLOR_MUTED)
        _format_cell_text(row.cells[1], value, font_size=10, color=COLOR_DARK)
        row.cells[0].width = Inches(2)
        row.cells[1].width = Inches(3)

    doc.add_page_break()


def _add_toc_placeholder(doc):
    h = doc.add_heading("Table of Contents", level=1)
    h.runs[0].font.color.rgb = BRAND_NAVY

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("(Update this field in Word: right-click and select 'Update Field' to generate the table of contents)")
    run.font.size = Pt(9)
    run.font.color.rgb = COLOR_MUTED
    run.italic = True

    p2 = doc.add_paragraph()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run1 = p2.add_run()
    run1._r.append(fldChar1)

    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
    run2 = p2.add_run()
    run2._r.append(instrText)

    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3 = p2.add_run()
    run3._r.append(fldChar2)

    run4 = p2.add_run("Right-click here and select 'Update Field' to generate TOC")
    run4.font.size = Pt(10)
    run4.font.color.rgb = COLOR_MUTED

    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5 = p2.add_run()
    run5._r.append(fldChar3)

    doc.add_page_break()


def _add_header_footer(doc, solution_name: str):
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = hp.add_run(f"{solution_name}  |  Technical Design Document")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = COLOR_MUTED

    hp_right = hp.add_run(f"\t\t{datetime.now().strftime('%B %Y')}")
    hp_right.font.name = "Calibri"
    hp_right.font.size = Pt(8)
    hp_right.font.color.rgb = COLOR_MUTED

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = fp.add_run("Confidential  |  Page ")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = COLOR_MUTED

    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    r1 = fp.add_run()
    r1._r.append(fldChar1)
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    r2 = fp.add_run()
    r2._r.append(instrText)
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    r3 = fp.add_run()
    r3._r.append(fldChar2)


def _add_styled_table(doc, rows: list[str]):
    parsed = []
    for row in rows:
        cells = [c.strip() for c in row.strip("|").split("|")]
        if all(re.match(r"^[-:]+$", c) for c in cells):
            continue
        parsed.append(cells)

    if len(parsed) < 1:
        return

    cols = len(parsed[0])
    table = doc.add_table(rows=len(parsed), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="D0D0D0"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="D0D0D0"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="D0D0D0"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="D0D0D0"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="D0D0D0"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="D0D0D0"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    for i, row_data in enumerate(parsed):
        for j, cell_text in enumerate(row_data):
            if j < cols:
                cell = table.rows[i].cells[j]
                clean = _clean_inline_md(cell_text)

                if i == 0:
                    _set_cell_shading(cell, "1B2A4A")
                    _format_cell_text(cell, clean, font_size=9, bold=True, color=COLOR_WHITE)
                else:
                    bg = "F8F9FA" if i % 2 == 0 else "FFFFFF"
                    _set_cell_shading(cell, bg)
                    _format_cell_text(cell, clean, font_size=9, color=COLOR_BODY)

                _set_cell_margins(cell)

    doc.add_paragraph("").paragraph_format.space_after = Pt(4)


def export_to_docx(sections: list[dict], solution_name: str) -> bytes:
    doc = Document()

    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    _setup_styles(doc)
    _add_cover_page(doc, solution_name)
    _add_toc_placeholder(doc)
    _add_header_footer(doc, solution_name)

    sorted_sections = sorted(sections, key=lambda s: s.get("order", 0))

    for section_data in sorted_sections:
        content = section_data.get("content", "")
        blocks = _strip_markdown(content)
        table_rows = []

        for block in blocks:
            btype = block["type"]
            text = _clean_inline_md(block["text"])

            if btype == "blank":
                if table_rows:
                    _add_styled_table(doc, table_rows)
                    table_rows = []
                continue
            elif btype == "table_row":
                table_rows.append(block["text"])
                continue
            else:
                if table_rows:
                    _add_styled_table(doc, table_rows)
                    table_rows = []

            if btype == "h1":
                doc.add_heading(text, level=1)
            elif btype == "h2":
                doc.add_heading(text, level=2)
            elif btype == "h3":
                doc.add_heading(text, level=3)
            elif btype == "h4":
                h = doc.add_heading(text, level=4)
                h.runs[0].font.size = Pt(11)
            elif btype in ("h5", "h6"):
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = COLOR_DARK
            elif btype == "bullet":
                p = doc.add_paragraph(style="List Bullet")
                p.clear()
                run = p.add_run(text)
                run.font.size = Pt(10)
                run.font.color.rgb = COLOR_BODY
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
            elif btype == "numbered":
                p = doc.add_paragraph(style="List Number")
                p.clear()
                run = p.add_run(text)
                run.font.size = Pt(10)
                run.font.color.rgb = COLOR_BODY
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
            elif btype == "hr":
                p = doc.add_paragraph()
                run = p.add_run("━" * 60)
                run.font.color.rgb = RGBColor(0xD0, 0xD0, 0xD0)
                run.font.size = Pt(6)
            elif btype == "code_block":
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(block["text"])
                run.font.name = "Consolas"
                run.font.size = Pt(8)
                run.font.color.rgb = COLOR_DARK
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
                p._p.get_or_add_pPr().append(shading)
            else:
                p = doc.add_paragraph(text)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(4)

        if table_rows:
            _add_styled_table(doc, table_rows)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _sanitize_for_pdf(text: str) -> str:
    replacements = {
        "\u2022": "-",
        "\u2013": "-",
        "\u2014": "--",
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\u2026": "...",
        "\u00a0": " ",
        "\u2192": "->",
        "\u2190": "<-",
        "\u2193": "v",
        "\u2191": "^",
        "\u2501": "-",
    }
    for char, repl in replacements.items():
        text = text.replace(char, repl)
    try:
        text.encode("latin-1")
    except UnicodeEncodeError:
        text = text.encode("latin-1", errors="replace").decode("latin-1")
    return text


class PDFDoc(FPDF):
    def __init__(self, solution_name: str):
        super().__init__()
        self.solution_name = solution_name
        self.set_auto_page_break(auto=True, margin=20)

    def header(self):
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(150, 150, 150)
        self.cell(0, 10, _sanitize_for_pdf(self.solution_name), align="L")
        self.ln(5)

    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(150, 150, 150)
        self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", align="C")


def export_to_pdf(sections: list[dict], solution_name: str) -> bytes:
    pdf = PDFDoc(solution_name)
    pdf.alias_nb_pages()
    pdf.add_page()

    pdf.set_font("Helvetica", "B", 20)
    pdf.set_text_color(30, 30, 30)
    pdf.cell(0, 15, _sanitize_for_pdf(solution_name), ln=True, align="C")
    pdf.set_font("Helvetica", "", 12)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 10, "Technical Design Document", ln=True, align="C")
    pdf.ln(10)

    for section in sorted(sections, key=lambda s: s.get("order", 0)):
        content = section.get("content", "")
        blocks = _strip_markdown(content)
        table_rows = []

        for block in blocks:
            btype = block["type"]
            text = _sanitize_for_pdf(_clean_inline_md(block["text"]))

            if btype == "table_row":
                table_rows.append(text)
                continue
            else:
                if table_rows:
                    _render_pdf_table(pdf, table_rows)
                    table_rows = []

            pdf.set_x(pdf.l_margin)

            if btype == "blank":
                pdf.ln(3)
                continue

            if btype == "h1":
                pdf.ln(5)
                pdf.set_font("Helvetica", "B", 18)
                pdf.set_text_color(27, 42, 74)
                pdf.multi_cell(0, 9, text)
                pdf.ln(3)
            elif btype == "h2":
                pdf.ln(4)
                pdf.set_font("Helvetica", "B", 15)
                pdf.set_text_color(43, 87, 154)
                pdf.multi_cell(0, 8, text)
                pdf.ln(2)
            elif btype == "h3":
                pdf.ln(3)
                pdf.set_font("Helvetica", "B", 13)
                pdf.set_text_color(43, 87, 154)
                pdf.multi_cell(0, 7, text)
                pdf.ln(2)
            elif btype in ("h4", "h5", "h6"):
                pdf.set_font("Helvetica", "B", 11)
                pdf.set_text_color(45, 45, 45)
                pdf.multi_cell(0, 6, text)
                pdf.ln(1)
            elif btype == "bullet":
                pdf.set_font("Helvetica", "", 10)
                pdf.set_text_color(58, 58, 58)
                pdf.multi_cell(0, 5, f"  - {text}")
            elif btype == "numbered":
                pdf.set_font("Helvetica", "", 10)
                pdf.set_text_color(58, 58, 58)
                pdf.multi_cell(0, 5, text)
            elif btype == "hr":
                pdf.ln(2)
                y = pdf.get_y()
                pdf.set_draw_color(200, 200, 200)
                pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
                pdf.ln(2)
            else:
                pdf.set_font("Helvetica", "", 10)
                pdf.set_text_color(58, 58, 58)
                pdf.multi_cell(0, 5, text)
                pdf.ln(1)

        if table_rows:
            _render_pdf_table(pdf, table_rows)

    return bytes(pdf.output())


def _render_pdf_table(pdf: FPDF, rows: list[str]):
    parsed = []
    for row in rows:
        cells = [c.strip() for c in row.strip("|").split("|")]
        if all(re.match(r"^[-:]+$", c) for c in cells):
            continue
        parsed.append(cells)

    if not parsed:
        return

    cols = len(parsed[0])
    page_width = pdf.w - pdf.l_margin - pdf.r_margin
    col_width = page_width / cols

    for i, row_data in enumerate(parsed):
        if i == 0:
            pdf.set_font("Helvetica", "B", 9)
            pdf.set_fill_color(27, 42, 74)
            pdf.set_text_color(255, 255, 255)
        else:
            pdf.set_font("Helvetica", "", 9)
            if i % 2 == 0:
                pdf.set_fill_color(248, 249, 250)
            else:
                pdf.set_fill_color(255, 255, 255)
            pdf.set_text_color(58, 58, 58)

        for j, cell_text in enumerate(row_data):
            if j < cols:
                clean = _sanitize_for_pdf(_clean_inline_md(cell_text))
                pdf.cell(col_width, 6, clean[:50], border=1, fill=True)
        pdf.ln()

    pdf.set_text_color(58, 58, 58)
    pdf.ln(3)
